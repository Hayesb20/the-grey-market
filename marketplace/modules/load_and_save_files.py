# Brian Hayes
#16 Nov 2022

import sys
sys.path.append("marketplace/class_files")


import ast
import docx
from docx.shared import Inches
import os


from atv_class import Atv
from img_downloader import get_img_from_web

os_string = "C:/Users/hayesb3/Geany 1.38/Program Files/marketplace/the-grey-market/"

dict_of_filepaths = {
						"filepath_to_vehicle_database.txt" 			: os_string+"marketplace/txt_files/vehicle_database.txt",
						"filepath_to_vehicle_database_backup.txt"	: os_string+"marketplace/txt_files/vehicle_database_backup.txt"	
					}

# TESTED - given a file name and a database, will figure out which loading method to call
def load_file(filepath):		
	
	if filepath == dict_of_filepaths["filepath_to_vehicle_database.txt"]:
		database = load_file_helper(filepath)
		
	elif filepath == dict_of_filepaths["filepath_to_vehicle_database_backup.txt"]:
		database = load_file_helper(filepath)
	return database
			
# TESTED - Takes an empty list and filename, will load the list with contence of file (atvs)	
def load_file_helper(file):
	try: 
		with open(file, "r") as file_object:
			contence = file_object.read()
	except: print("invalid filename")
		
	database = []
	a_word = ""
	while contence != "":
		if contence[0] != "\n":
			a_word = a_word + contence[0]
			contence = contence[1:]
	
		elif contence[0] == "\n":	
			try: 
				kwargs = ast.literal_eval(a_word)
				if kwargs["classification"].lower() ==  "all terain vehicle" or kwargs["classification"].lower() == "four wheeler":
					newVehicle = Atv(**kwargs)
					database.append(newVehicle)	
					a_word = ""
					contence = contence[1:]	
				else: contence = contence[1:]	
			except: "Object creation failed"
		else: contence = contence[1:]
	return database
			
# TESTED - Will use a filename to determin which sub function needs to be called		
def save_file(database, filepath):
	if filepath == dict_of_filepaths["filepath_to_vehicle_database.txt"]:
		with open(filepath, "w") as file_object:
			long_string = convert_objects_in_vehicle_database_to_a_string(database)
			file_object.write(long_string)		
			
	elif filepath == dict_of_filepaths["filepath_to_vehicle_database_backup.txt"]:
		with open(filepath, "w") as file_object:
			long_string = convert_objects_in_vehicle_database_to_a_string(database)
			file_object.write(long_string)	

	elif filepath == "txt_files/test_txts/test_file_1.txt":
		with open(filepath, "w") as file_object:
			long_string = convert_objects_in_vehicle_database_to_a_string(database)
			file_object.write(long_string)	
	
	elif "Search Results" in filepath:
		# Opened a specific file or creates one if it doesnt exist, for writing only
		with open(filepath, "w+") as file_object: 
			list_of_strings = convert_list_to_string(database)
			create_docx(list_of_strings, get_img_from_web(database), filepath)

# TESTED - Given the atv_database list, returns a string
def convert_objects_in_vehicle_database_to_a_string(list_of_objects):
	string = ""
	for thing in list_of_objects:
		string = string + str(thing.get_essence_as_dict()) + "\n"
	return string

def create_docx(list_of_strings, list_of_imgs, filepath):

	doc = docx.Document()	# Create a docx object
	doc.add_heading("Search Results", 0) # Add a heading to the document
	for string, img in zip(list_of_strings, list_of_imgs):	# Going through two lists simultainusuly
		doc.add_picture(img, width=Inches(2), height=Inches(2))	# Adding an image to the documents
		doc.add_paragraph(string)	# Adding the text to the document
		os.remove(img)	# Deleting the downloaded images because we have them in the document now
	doc.save(filepath)	# Saving the file in the designated location

def convert_list_to_string(list):
	list_of_strings = []
	string = ""
	for tup in list:	# For each tuple in the list
		for char in tup[0]:	# For each character in the first element of the tuple
			if char == "\n": string = string + "\n"
			else: string = string + char
		list_of_strings.append(string)	# Adding the completed string to a list
		string = ""	# Reseting the string
	return list_of_strings



# NOT TESTED - Returns a string that contains information about an ATV to be orinted		
#needs to actually return a string to be printed. not print itself
def show_database(database):
	for item in database:
		print(item.get_essence_as_dict())
		
			
